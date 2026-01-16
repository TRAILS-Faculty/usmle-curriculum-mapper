import io
import pandas as pd
from docx import Document

def generate_reports(results_df: pd.DataFrame, summary_df: pd.DataFrame):
    # Word report
    doc = Document()
    doc.add_heading("USMLE Curriculum Mapping Report", level=1)
    doc.add_paragraph("Coverage rubric: Extensive (≥10 or in objectives), Moderate (3–9), Low (1–2).")
    for _, row in summary_df.iterrows():
        doc.add_heading(str(row["expanded_keywords"]), level=2)
        doc.add_paragraph(f"Mentions: {int(row['Mentions'])}")
        doc.add_paragraph(f"Coverage: {row['Coverage']}")
    word_buf = io.BytesIO()
    doc.save(word_buf)
    word_buf.seek(0)

    # Excel report
    excel_buf = io.BytesIO()
    with pd.ExcelWriter(excel_buf, engine="openpyxl") as w:
        results_df.to_excel(w, index=False, sheet_name="Sentences")
        summary_df.to_excel(w, index=False, sheet_name="Summary")
    excel_buf.seek(0)

    return word_buf, excel_buf
